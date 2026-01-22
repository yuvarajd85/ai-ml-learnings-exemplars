"""
LDDashRagChatbot
================

A Dash-based chatbot UI backed by Google Gemini (via LangChain) with optional
Retrieval-Augmented Generation (RAG) over an uploaded document.

The application supports uploading common document types (txt/md/csv/json/pdf/docx),
building a per-session vector index (Chroma), and answering questions using a
LangChain Expression Language (LCEL) RAG chain.

.. warning::
   This implementation stores vector indexes in a process-level in-memory dict
   (``RAG_SESSIONS``). This is suitable for local development and demos, but is
   not safe for multi-worker deployments (e.g., Gunicorn with multiple workers)
   without persistence or an external store.

Environment Variables
---------------------
- ``GOOGLE_API_KEY`` (required): Google Gemini API key.
- ``GEMINI_MODEL`` (optional): Gemini chat model name. Default: ``gemini-2.5-flash-lite``.
- ``EMBEDDING_MODEL`` (optional): Google embedding model. Default: ``models/text-embedding-004``.
- ``TEMPERATURE`` (optional): LLM temperature. Default: ``0.2``.
- ``CHUNK_SIZE`` (optional): Chunk size for text splitting. Default: ``1000``.
- ``CHUNK_OVERLAP`` (optional): Chunk overlap for text splitting. Default: ``150``.
- ``TOP_K`` (optional): Number of retrieved chunks for RAG. Default: ``5``.
- ``MAX_TURNS`` (optional): Rolling window of chat turns kept in UI memory. Default: ``12``.

Running
-------
Run locally:

.. code-block:: bash

   python LDDashRagChatbot.py

Then open:

- http://127.0.0.1:8050

Created on 1/20/26 at 9:24 PM
By yuvarajdurairaj
Module Name LDDashRagChatbot
"""

import os
import base64
import json
import uuid
from typing import List

from dotenv import load_dotenv
from dash import Dash, html, dcc, Input, Output, State, callback_context, no_update

from pypdf import PdfReader
from docx import Document as DocxDocument

from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from langchain_chroma import Chroma
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser

load_dotenv()

# -----------------------------
# Config
# -----------------------------
IS_SPHINX = os.getenv("SPHINX_BUILD") == "1"

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY and not IS_SPHINX:
    raise RuntimeError("Missing GOOGLE_API_KEY in env/.env")

GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash-lite")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "models/text-embedding-004")

TEMPERATURE = float(os.getenv("TEMPERATURE", "0.2"))
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "1000"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "150"))
TOP_K = int(os.getenv("TOP_K", "5"))

MAX_TURNS = int(os.getenv("MAX_TURNS", "12"))  # chat history window

# -----------------------------
# LLM + Embeddings
# -----------------------------
llm = ChatGoogleGenerativeAI(
    model=GEMINI_MODEL,
    temperature=TEMPERATURE,
    google_api_key=GOOGLE_API_KEY,
)

embeddings = GoogleGenerativeAIEmbeddings(
    model=EMBEDDING_MODEL,
    google_api_key=GOOGLE_API_KEY,
)

# -----------------------------
# RAG prompt + chain builder (LCEL)
# -----------------------------
SYSTEM_PROMPT = (
    "You are a precise senior engineering mentor. "
    "Use the provided context to answer. "
    "If the context does not contain the answer, say so explicitly and then answer from general knowledge "
    "while clearly labeling it as 'General knowledge' vs 'From document'. "
    "Keep it concise and correct."
)

prompt = ChatPromptTemplate.from_messages(
    [
        ("system", SYSTEM_PROMPT),
        (
            "human",
            "Question:\n{question}\n\n"
            "Context:\n{context}\n\n"
            "Answer:",
        ),
    ]
)


def format_docs(docs: List[Document]) -> str:
    """
    Format retrieved documents into a single context string.

    This function is intended for use in an LCEL chain as a post-processing step
    after retrieval. It includes basic source metadata when available.

    Parameters
    ----------
    docs:
        A list of LangChain ``Document`` objects returned by a retriever.

    Returns
    -------
    str
        A formatted context string suitable for insertion into a prompt. If no
        documents are provided, returns ``"NO_RELEVANT_CONTEXT"``.
    """
    if not docs:
        return "NO_RELEVANT_CONTEXT"

    chunks: List[str] = []
    for i, d in enumerate(docs, start=1):
        src = d.metadata.get("source", "unknown")
        page = d.metadata.get("page", None)
        loc = f"{src}" + (f" (page {page})" if page is not None else "")
        chunks.append(f"[{i}] {loc}\n{d.page_content}")

    return "\n\n---\n\n".join(chunks)


def build_rag_chain(retriever):
    """
    Build an LCEL Retrieval-Augmented Generation (RAG) chain.

    The chain follows this conceptual structure:

    - Retrieve relevant documents for a question
    - Format documents into a promptable context string
    - Combine ``{context}`` and ``{question}`` into a chat prompt
    - Invoke the Gemini chat model
    - Parse the model output into a plain string

    LCEL equivalent:

    .. code-block:: python

       (
         {"context": retriever | format_docs, "question": RunnablePassthrough()}
         | prompt
         | llm
         | StrOutputParser()
       )

    Parameters
    ----------
    retriever:
        A LangChain retriever implementing the Runnable interface and returning
        a list of ``Document`` objects.

    Returns
    -------
    Runnable
        An LCEL runnable chain that accepts a question string and returns an
        answer string.
    """
    return (
        {
            "context": retriever | format_docs,
            "question": RunnablePassthrough(),
        }
        | prompt
        | llm
        | StrOutputParser()
    )


# -----------------------------
# Upload decoding + parsing
# -----------------------------
def decode_upload(contents: str) -> bytes:
    """
    Decode Dash ``dcc.Upload`` contents into raw bytes.

    Dash upload contents are provided as a data URL:

    ``data:<mime>;base64,<payload>``

    Parameters
    ----------
    contents:
        The upload contents string from ``dcc.Upload(contents=...)``.

    Returns
    -------
    bytes
        The decoded file contents.
    """
    _header, b64data = contents.split(",", 1)
    return base64.b64decode(b64data)


def parse_txt_like(raw: bytes) -> str:
    """
    Decode bytes into text for plain-text-like files.

    Attempts UTF-8 first and falls back to Latin-1 with replacement.

    Parameters
    ----------
    raw:
        Raw bytes for a text-like file.

    Returns
    -------
    str
        Decoded text content.
    """
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("latin-1", errors="replace")


def parse_json(raw: bytes) -> str:
    """
    Parse JSON bytes into a pretty-printed text representation.

    If JSON parsing fails, falls back to text decoding.

    Parameters
    ----------
    raw:
        Raw bytes for a JSON file.

    Returns
    -------
    str
        Pretty-printed JSON string or decoded text fallback.
    """
    try:
        obj = json.loads(raw.decode("utf-8"))
        return json.dumps(obj, indent=2, ensure_ascii=False)
    except Exception:
        return parse_txt_like(raw)


def parse_pdf(raw: bytes, filename: str) -> List[Document]:
    """
    Extract text from a PDF and return it as a list of Documents.

    Each PDF page becomes a separate ``Document``. Pages with no extractable
    text are skipped.

    Parameters
    ----------
    raw:
        Raw bytes of the PDF file.
    filename:
        The original filename, stored in ``Document.metadata['source']``.

    Returns
    -------
    list[Document]
        One ``Document`` per page containing extractable text.
    """
    import io

    reader = PdfReader(io.BytesIO(raw))
    docs: List[Document] = []

    for idx, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if text:
            docs.append(
                Document(
                    page_content=text,
                    metadata={"source": filename, "page": idx + 1},
                )
            )

    return docs


def parse_docx(raw: bytes, filename: str) -> List[Document]:
    """
    Extract text from a DOCX and return it as a list of Documents.

    The DOCX content is extracted from paragraphs, joined, and returned as a
    single ``Document`` (which will later be chunked by the text splitter).

    Parameters
    ----------
    raw:
        Raw bytes of the DOCX file.
    filename:
        The original filename, stored in ``Document.metadata['source']``.

    Returns
    -------
    list[Document]
        A single ``Document`` containing extracted DOCX text, or an empty list
        if no text was found.
    """
    import io

    doc = DocxDocument(io.BytesIO(raw))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    text = "\n".join(paragraphs).strip()

    if not text:
        return []

    return [Document(page_content=text, metadata={"source": filename})]


def load_documents_from_upload(contents: str, filename: str) -> List[Document]:
    """
    Load an uploaded file into LangChain Documents based on extension.

    Supported extensions:
    - ``txt``, ``md``, ``csv`` -> single Document
    - ``json`` -> single Document (pretty-printed)
    - ``pdf`` -> one Document per page
    - ``docx`` -> single Document

    Parameters
    ----------
    contents:
        The Dash upload contents string from ``dcc.Upload``.
    filename:
        The uploaded filename (used to infer file type and stored as metadata).

    Returns
    -------
    list[Document]
        Extracted Documents representing the file content.

    Raises
    ------
    ValueError
        If the file extension is unsupported.
    """
    raw = decode_upload(contents)
    ext = filename.lower().split(".")[-1] if "." in filename else ""

    if ext in ("txt", "md", "csv"):
        return [Document(page_content=parse_txt_like(raw), metadata={"source": filename})]
    if ext == "json":
        return [Document(page_content=parse_json(raw), metadata={"source": filename})]
    if ext == "pdf":
        return parse_pdf(raw, filename)
    if ext == "docx":
        return parse_docx(raw, filename)

    raise ValueError(f"Unsupported file type: .{ext}. Supported: txt/md/csv/json/pdf/docx")


# -----------------------------
# Chunk + vectorstore build
# -----------------------------
splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP,
)


def build_vectorstore(docs: List[Document], collection_name: str) -> Chroma:
    """
    Build a Chroma vector store from Documents.

    Documents are chunked using the module-level ``splitter`` and embedded using
    the module-level Gemini embedding model.

    Parameters
    ----------
    docs:
        Input Documents to index.
    collection_name:
        Name of the Chroma collection.

    Returns
    -------
    Chroma
        A Chroma vector store containing the embedded chunks.
    """
    chunks = splitter.split_documents(docs)
    vs = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        collection_name=collection_name,
    )
    return vs


# -----------------------------
# Simple in-memory per-session storage (DEV ONLY)
# -----------------------------
#: In-memory store mapping a Dash session id to its RAG index state.
#:
#: Keys are session ids (strings). Values are dicts with at least:
#:
#: - ``vs``: Chroma vector store
#: - ``retriever``: Retriever derived from the vector store
#: - ``filename``: Uploaded filename
#: - ``doc_count``: Number of source Documents produced by parsing
RAG_SESSIONS = {}


def trim_history(history):
    """
    Trim chat history to a rolling window of recent turns.

    The UI stores messages as dicts of the form ``{"role": "...", "content": "..."}``.
    This function keeps only the last ``MAX_TURNS * 2`` entries (user+assistant).

    Parameters
    ----------
    history:
        List of chat message dicts.

    Returns
    -------
    list
        Trimmed history list.
    """
    if not history:
        return []
    return history[-(MAX_TURNS * 2):]


# -----------------------------
# Dash UI
# -----------------------------
app = Dash(__name__)
server = app.server

app.layout = html.Div(
    style={
        "maxWidth": "900px",
        "margin": "0 auto",
        "padding": "16px",
        "fontFamily": "system-ui, -apple-system, Segoe UI, Roboto, sans-serif",
    },
    children=[
        html.H2("Dash Chatbot (Gemini + LangChain RAG)", style={"marginBottom": "8px"}),

        dcc.Store(id="session-id", data=str(uuid.uuid4())),
        dcc.Store(id="chat-store", data=[]),

        html.Div(
            id="chat-window",
            style={
                "height": "520px",
                "overflowY": "auto",
                "border": "1px solid #ddd",
                "borderRadius": "12px",
                "padding": "12px",
                "background": "#fafafa",
            },
        ),

        html.Div(style={"height": "12px"}),

        html.Div(
            style={"display": "flex", "gap": "10px", "alignItems": "center"},
            children=[
                dcc.Upload(
                    id="upload-file",
                    children=html.Button(
                        "Attach (build RAG index)",
                        style={
                            "borderRadius": "12px",
                            "border": "1px solid #ddd",
                            "background": "white",
                            "cursor": "pointer",
                            "padding": "8px 12px",
                            "fontWeight": "600",
                        },
                    ),
                    multiple=False,
                ),
                html.Button(
                    "Clear attachment",
                    id="clear-attachment-btn",
                    n_clicks=0,
                    style={
                        "borderRadius": "12px",
                        "border": "1px solid #ddd",
                        "background": "white",
                        "cursor": "pointer",
                        "padding": "8px 12px",
                    },
                ),
                html.Div(id="attachment-status", style={"color": "#666"}, children="No attachment indexed."),
            ],
        ),

        html.Div(style={"height": "10px"}),

        html.Div(
            style={"display": "flex", "gap": "8px"},
            children=[
                dcc.Input(
                    id="user-input",
                    type="text",
                    placeholder="Type a message… (Enter to send)",
                    debounce=True,
                    style={
                        "flex": "1",
                        "height": "44px",
                        "borderRadius": "12px",
                        "padding": "0 12px",
                        "border": "1px solid #ddd",
                    },
                ),
                html.Button(
                    "Send",
                    id="send-btn",
                    n_clicks=0,
                    style={
                        "width": "110px",
                        "borderRadius": "12px",
                        "border": "1px solid #ddd",
                        "background": "white",
                        "cursor": "pointer",
                        "fontWeight": "600",
                    },
                ),
            ],
        ),

        html.Div(style={"height": "10px"}),

        html.Div(
            style={"display": "flex", "gap": "8px"},
            children=[
                html.Button(
                    "Clear chat",
                    id="clear-btn",
                    n_clicks=0,
                    style={
                        "borderRadius": "12px",
                        "border": "1px solid #ddd",
                        "background": "white",
                        "cursor": "pointer",
                    },
                ),
                html.Div(
                    style={"color": "#666", "paddingTop": "6px"},
                    children=f"Model: {GEMINI_MODEL} | Embeddings: {EMBEDDING_MODEL} | top_k={TOP_K} | chunk={CHUNK_SIZE}/{CHUNK_OVERLAP}",
                ),
            ],
        ),
    ],
)


def render_chat(history):
    """
    Render chat messages as simple "bubble" components.

    Parameters
    ----------
    history:
        List of chat message dicts with keys ``role`` and ``content``.

    Returns
    -------
    list
        A list of Dash HTML components representing the chat history.
    """
    bubbles = []
    for msg in history:
        role = msg.get("role", "assistant")
        content = msg.get("content", "")
        is_user = role == "user"

        bubbles.append(
            html.Div(
                style={
                    "display": "flex",
                    "justifyContent": "flex-end" if is_user else "flex-start",
                    "marginBottom": "10px",
                },
                children=[
                    html.Div(
                        content,
                        style={
                            "maxWidth": "80%",
                            "whiteSpace": "pre-wrap",
                            "padding": "10px 12px",
                            "borderRadius": "12px",
                            "border": "1px solid #ddd",
                            "background": "#e8f0fe" if is_user else "white",
                        },
                    )
                ],
            )
        )

    if not bubbles:
        bubbles = [html.Div("Upload a doc (optional) and ask a question.", style={"color": "#888"})]

    return bubbles


# -----------------------------
# Callbacks: Upload -> Build RAG index
# -----------------------------
@app.callback(
    Output("attachment-status", "children"),
    Input("upload-file", "contents"),
    State("upload-file", "filename"),
    State("session-id", "data"),
    prevent_initial_call=True,
)
def on_upload(contents, filename, session_id):
    """
    Dash callback: handle file upload and build the RAG vector index.

    Parameters
    ----------
    contents:
        Dash upload contents string (base64 data URL).
    filename:
        Uploaded filename.
    session_id:
        Per-browser session identifier.

    Returns
    -------
    str
        A status message describing index build success/failure.
    """
    if not contents or not filename:
        return no_update

    try:
        docs = load_documents_from_upload(contents, filename)
        if not docs:
            return f"[Attachment error] No extractable text found in {filename}"

        collection_name = f"rag_{session_id}"
        vs = build_vectorstore(docs, collection_name=collection_name)
        retriever = vs.as_retriever(search_kwargs={"k": TOP_K})

        RAG_SESSIONS[session_id] = {
            "vs": vs,
            "retriever": retriever,
            "filename": filename,
            "doc_count": len(docs),
        }

        # Quick stats: number of indexed chunks (internal API; fine for dev)
        chunk_count = vs._collection.count()
        return f"Indexed: {filename} | pages/docs={len(docs)} | chunks={chunk_count} | top_k={TOP_K}"

    except Exception as e:
        return f"[Attachment error] {type(e).__name__}: {e}"


@app.callback(
    Output("attachment-status", "children", allow_duplicate=True),
    Input("clear-attachment-btn", "n_clicks"),
    State("session-id", "data"),
    prevent_initial_call=True,
)
def clear_attachment(n, session_id):
    """
    Dash callback: clear the RAG index for the current session.

    Parameters
    ----------
    n:
        Number of clicks on the "Clear attachment" button.
    session_id:
        Per-browser session identifier.

    Returns
    -------
    str
        A status message indicating no attachment is indexed.
    """
    if not n:
        return no_update

    if session_id in RAG_SESSIONS:
        del RAG_SESSIONS[session_id]

    return "No attachment indexed."


# -----------------------------
# Callbacks: Chat
# -----------------------------
@app.callback(
    Output("chat-store", "data"),
    Output("user-input", "value"),
    Input("send-btn", "n_clicks"),
    Input("clear-btn", "n_clicks"),
    Input("user-input", "n_submit"),
    State("user-input", "value"),
    State("chat-store", "data"),
    State("session-id", "data"),
    prevent_initial_call=True,
)
def on_send_or_clear(send_clicks, clear_clicks, n_submit, user_text, history, session_id):
    """
    Dash callback: send a user message (or clear chat) and return updated history.

    If a RAG index exists for the current session, the callback uses the LCEL RAG
    chain to answer the user question. Otherwise it falls back to a plain LLM call.

    Parameters
    ----------
    send_clicks:
        Number of clicks on the "Send" button.
    clear_clicks:
        Number of clicks on the "Clear chat" button.
    n_submit:
        ``dcc.Input`` submit count (triggered when pressing Enter).
    user_text:
        Current text in the user input box.
    history:
        Chat history stored in ``dcc.Store``.
    session_id:
        Per-browser session identifier.

    Returns
    -------
    tuple[list, str]
        Updated chat history and the cleared input box value.
    """
    ctx = callback_context
    if not ctx.triggered:
        return no_update, no_update

    triggered_id = ctx.triggered[0]["prop_id"].split(".")[0]

    if triggered_id == "clear-btn":
        return [], ""

    if not user_text or not user_text.strip():
        return no_update, no_update

    history = history or []
    user_text = user_text.strip()

    history.append({"role": "user", "content": user_text})
    history = trim_history(history)

    try:
        rag_state = RAG_SESSIONS.get(session_id)

        if rag_state and rag_state.get("retriever"):
            rag_chain = build_rag_chain(rag_state["retriever"])
            assistant_text = rag_chain.invoke(user_text)
        else:
            assistant_text = llm.invoke(user_text).content

        assistant_text = (assistant_text or "").strip() or "[No response]"
    except Exception as e:
        assistant_text = f"[ERROR] {type(e).__name__}: {e}"

    history.append({"role": "assistant", "content": assistant_text})
    history = trim_history(history)

    return history, ""


@app.callback(
    Output("chat-window", "children"),
    Input("chat-store", "data"),
)
def update_chat_window(history):
    """
    Dash callback: render the chat window from stored history.

    Parameters
    ----------
    history:
        Chat history stored in ``dcc.Store``.

    Returns
    -------
    list
        Rendered chat components.
    """
    return render_chat(history or [])


if __name__ == "__main__":
    # Prefer run_server() for Dash apps. app.run() is not consistently supported.
    app.run(debug=True)